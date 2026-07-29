"""DOCX report generation.

Two report types for v1: **inventory snapshot** and **update-job report**, in a
clean tabular layout suitable for inclusion in the Netweb/Tyrone PASS/FAIL
reporting workflow. Templates can be refined as the house style demands.

``python-docx`` is imported lazily so the app launches even when it isn't
installed — the user only sees the import error when they actually click
Export, with a clear "pip install python-docx" message.
"""
from __future__ import annotations

import logging
from datetime import datetime
from pathlib import Path

log = logging.getLogger(__name__)


def _docx():
    try:
        from docx import Document  # type: ignore
        return Document
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError(
            "python-docx is not installed. Run: pip install python-docx"
        ) from exc


def _now_str() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def render_inventory_snapshot(servers: list[dict], output_path: Path) -> Path:
    Document = _docx()
    doc = Document()

    doc.add_heading("ClusterOne — Inventory Snapshot", level=0)
    meta = doc.add_paragraph()
    meta.add_run(f"Generated: {_now_str()}\n").italic = True
    meta.add_run(f"Total servers: {len(servers)}").bold = True

    table = doc.add_table(rows=1, cols=7)
    table.style = "Light Grid"
    header = ["IP", "Hostname", "OEM", "Model", "BMC", "BIOS", "Status"]
    for i, h in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.font.bold = True

    for s in servers:
        row = table.add_row().cells
        row[0].text = s.get("ip") or ""
        row[1].text = s.get("hostname") or ""
        row[2].text = (s.get("oem") or "").capitalize()
        row[3].text = s.get("model") or ""
        row[4].text = s.get("bmc_version") or ""
        row[5].text = s.get("bios_version") or ""
        row[6].text = (s.get("status") or "").upper()

    doc.add_paragraph()
    sig = doc.add_paragraph()
    sig.add_run("Validated by: ____________________      Date: ____________________").italic = True

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    log.info("Inventory report written: %s", output_path)
    return output_path


def render_update_job_report(job: dict, tasks: list[dict], output_path: Path) -> Path:
    Document = _docx()
    doc = Document()

    doc.add_heading(f"ClusterOne — Update Job #{job.get('id', '?')}", level=0)

    summary = doc.add_paragraph()
    summary.add_run(f"Generated: {_now_str()}\n").italic = True
    summary.add_run(f"Firmware type: {job.get('fw_type', '?')}\n").bold = True
    summary.add_run(f"State: {(job.get('state') or '?').upper()}\n")
    summary.add_run(f"Created: {job.get('created_at', '?')}\n")
    summary.add_run(
        f"Total: {job.get('total', 0)}    "
        f"Succeeded: {job.get('succeeded', 0)}    "
        f"Failed: {job.get('failed', 0)}\n"
    )

    table = doc.add_table(rows=1, cols=6)
    table.style = "Light Grid"
    header = ["Server", "OEM", "Firmware", "Before → After", "Result", "Notes"]
    for i, h in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.font.bold = True

    for t in tasks:
        row = table.add_row().cells
        row[0].text = t.get("server_hostname") or t.get("server_ip") or "?"
        row[1].text = (t.get("server_oem") or "").capitalize()
        fw_type = t.get("firmware_type") or ""
        fw_ver = t.get("firmware_version") or ""
        row[2].text = f"{fw_type} {fw_ver}".strip()
        row[3].text = f"{t.get('version_before') or '?'} → {t.get('version_after') or '?'}"
        state = (t.get("state") or "").lower()
        row[4].text = "PASS" if state == "completed" else "FAIL" if state == "failed" else state.upper()
        row[5].text = t.get("message") or ""

    doc.add_paragraph()
    sig = doc.add_paragraph()
    sig.add_run("Engineer: ____________________      Reviewer: ____________________      Date: ____________________").italic = True

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    log.info("Update-job report written: %s", output_path)
    return output_path
